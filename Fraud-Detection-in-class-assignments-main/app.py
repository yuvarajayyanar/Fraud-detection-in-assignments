import torch
from flask import Flask, request, jsonify, render_template, redirect, url_for, render_template_string, make_response
import json
import logging
import numpy as np
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from transformers import AutoTokenizer, AutoModelForSequenceClassification, pipeline
import torch.nn.functional as F
import re
import math
from collections import Counter
import nltk
from nltk.corpus import stopwords
from docx import Document
import PyPDF2
import os
from googleapiclient.discovery import build
import requests
import pdfkit
from datetime import datetime
import sqlite3
from sklearn.preprocessing import normalize
from io import BytesIO
import base64

# Configure logging
logging.basicConfig(level=logging.INFO)

app = Flask(__name__)

# Configure upload folder
UPLOAD_FOLDER = 'uploads'
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# Download NLTK data
nltk.download('stopwords')
nltk.download('wordnet')

searchEngine_API = 'AIzaSyAJfOxvSEfonJE_l96vHdL_5ega41fgutc'
searchEngine_Id = 'd78348c105b614496'

# Render HTML Pages
@app.route('/')
def index():
    return render_template("index.html")

@app.route('/cluster/')
def cluster():
    return render_template("cluster.html")

@app.route('/ai_detection/')
def ai_detection():
    return render_template("ai_detection.html")

@app.route('/code_cluster/')
def code_cluster():
    return render_template("code_cluster.html")

@app.route('/assignment/')
def assignment():
    return render_template("assignment.html")

@app.route('/plagiarism_checker/')
def plagiarism_checker():
    return render_template("plagiarism_checker.html", link={}, percent=0)

@app.route('/home')
def home():
    return render_template("index.html", link={}, percent=0)
@app.route('/trend_analysis')
def trend_analysis():
    return render_template("trend_analysis.html")
# Function to create clusters based on cosine similarity
def create_clusters(texts, threshold=0.3):
    try:
        if len(texts) == 1:
            logging.info("Only one file provided, creating a single cluster without similarity.")
            return [[0]], []
        vectorizer = TfidfVectorizer()
        tfidf_matrix = vectorizer.fit_transform(texts)
        cosine_sim = cosine_similarity(tfidf_matrix)
        print("Cosine Similarity Matrix:")
        for row in cosine_sim:
            print([round(val, 3) for val in row])
        clusters = []
        visited = set()
        cluster_similarities = []
        for i in range(len(texts)):
            if i not in visited:
                current_cluster = [i]
                similarities = []
                for j in range(i + 1, len(texts)):
                    if cosine_sim[i][j] >= threshold:
                        current_cluster.append(j)
                        visited.add(j)
                        similarities.append(cosine_sim[i][j])
                avg_similarity = round(np.mean(similarities) * 100, 2) if similarities else 100
                clusters.append(current_cluster)
                cluster_similarities.append(avg_similarity)
                visited.add(i)
        return clusters, cluster_similarities
    except Exception as e:
        logging.error(f"Error in create_clusters: {str(e)}")
        return [], []

@app.route('/compare-texts/', methods=['POST'])
def cluster_files():
    try:
        data = request.get_json()
        files_data = data.get("files", [])
        texts = [file["content"] for file in files_data]
        names = [file["name"] for file in files_data]
        if not texts:
            return jsonify({'error': 'No valid files found'}), 400
        clusters, cluster_similarities = create_clusters(texts, threshold=0.6)
        result = []
        for idx, cluster in enumerate(clusters):
            cluster_data = {
                "cluster_id": idx + 1,
                "names": [names[i] for i in cluster],
            }
            if len(cluster) > 1:
                cluster_data["avg_similarity"] = cluster_similarities[idx]
            result.append(cluster_data)
        return jsonify(result)
    except Exception as e:
        logging.error(f"Error in cluster_files: {str(e)}")
        return jsonify({'error': 'Internal server error'}), 500

@app.route('/generate_text_cluster_pdf', methods=['POST'])
def generate_text_cluster_pdf():
    try:
        data = request.get_json()
        clusters = data.get('clusters', [])
        html = render_template_string('''
            <!DOCTYPE html>
            <html>
            <head>
                <style>
                    body { font-family: Arial, sans-serif; margin: 1in; line-height: 1.6; }
                    h1 { color: #2b2d42; font-size: 24px; border-bottom: 2px solid #4361ee; padding-bottom: 5px; }
                    .report-info { font-size: 12px; color: #666; margin-bottom: 20px; }
                    .cluster { margin: 20px 0; padding: 10px; background-color: #f9f9f9; border-radius: 5px; }
                    .cluster-header { display: flex; justify-content: space-between; align-items: center; margin-bottom: 10px; }
                    .cluster-title { font-size: 18px; font-weight: bold; color: #2b2d42; }
                    .similarity { color: #4361ee; font-weight: bold; }
                    ul { list-style-type: none; padding-left: 20px; }
                    li { margin: 5px 0; font-size: 12px; }
                    li:before { content: "• "; color: #4361ee; }
                </style>
            </head>
            <body>
                <h1>Text Document Clustering Report</h1>
                <div class="report-info">Generated on: {{ timestamp }}</div>
                {% if clusters %}
                    {% for cluster in clusters %}
                    <div class="cluster">
                        <div class="cluster-header">
                            <span class="cluster-title">Cluster {{ cluster.cluster_id }}</span>
                            {% if cluster.avg_similarity is defined %}
                                <span class="similarity">Similarity: {{ cluster.avg_similarity }}%</span>
                            {% endif %}
                        </div>
                        <ul>
                            {% for file in cluster.names %}
                            <li>{{ file }}</li>
                            {% endfor %}
                        </ul>
                    </div>
                    {% endfor %}
                {% else %}
                    <p>No clusters found.</p>
                {% endif %}
            </body>
            </html>
        ''', clusters=clusters, timestamp=datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
        try:
            config = pdfkit.configuration()
        except Exception as e:
            logging.warning(f"wkhtmltopdf not found in PATH, using hardcoded path: {str(e)}")
            config = pdfkit.configuration(wkhtmltopdf='C:/Program Files/wkhtmltopdf/bin/wkhtmltopdf.exe')
        pdf = pdfkit.from_string(html, False, options={
            'page-size': 'A4',
            'margin-top': '0.5in',
            'margin-right': '0.5in',
            'margin-bottom': '0.5in',
            'margin-left': '0.5in',
            'encoding': 'UTF-8'
        }, configuration=config)
        response = make_response(pdf)
        response.headers['Content-Type'] = 'application/pdf'
        response.headers['Content-Disposition'] = 'attachment; filename=text_cluster_report.pdf'
        return response
    except Exception as e:
        logging.error(f"PDF generation error: {str(e)}")
        return jsonify({'error': f'Failed to generate PDF: {str(e)}'}), 500

# AI Detection
tokenizer = AutoTokenizer.from_pretrained("pritamdeb68/BERTAIDetector")
model = AutoModelForSequenceClassification.from_pretrained("pritamdeb68/BERTAIDetector")
model.eval()

@app.route('/detect_ai_content/', methods=['POST'])
def detect_ai_content():
    try:
        data = request.get_json()
        text = data.get("text", "").strip()
        if not text:
            return jsonify({"error": "Empty text input"}), 400
        inputs = tokenizer(text, return_tensors="pt", truncation=True, padding=True, max_length=512)
        with torch.no_grad():
            outputs = model(**inputs)
            logits = outputs.logits
            probabilities = F.softmax(logits, dim=-1)[0].tolist()
        ai_probability = round(probabilities[1], 2)
        human_probability = round(probabilities[0], 2)
        return jsonify({
            "ai_probability": ai_probability,
            "human_probability": human_probability
        })
    except Exception as e:
        logging.error(f"Error in detect_ai_content_api: {str(e)}", exc_info=True)
        return jsonify({"error": "Internal server error"}), 500

# Restored original /ai_refined/ endpoint with RapidAPI
@app.route('/ai_refined/', methods=['POST'])
def ai_refined():
    try:
        # Get text from the frontend request
        data = request.get_json()
        text = data.get("text", "").strip()
        if not text:
            return jsonify({"error": "Empty text input"}), 400

        # RapidAPI AI Content Detector setup
        url = "https://ai-content-detector-ai-gpt.p.rapidapi.com/api/detectText/"
        headers = {
            "x-rapidapi-key": "f9f4afc3bbmshe706902778d4a95p17967fjsne94dbf9c3400",
            "x-rapidapi-host": "ai-content-detector-ai-gpt.p.rapidapi.com",
            "Content-Type": "application/json"
        }
        payload = {"text": text}

        # Send request to RapidAPI
        response = requests.post(url, json=payload, headers=headers)
        
        # Check if request was successful
        if response.status_code != 200:
            logging.error(f"RapidAPI error: {response.status_code} - {response.text}")
            return jsonify({"error": f"API request failed: {response.status_code}"}), 500
        
        # Print the raw RapidAPI response to the terminal
        print("RapidAPI Response:", response.json())

        # Parse RapidAPI response
        result = response.json()
        
        # Extract relevant fields
        fake_percentage = result.get("fakePercentage", 0.0) / 100  # Convert to 0-1 scale
        is_human = result.get("isHuman", 0.0) / 100  # Convert to 0-1 scale
        ai_score = fake_percentage  # Use fakePercentage as AI probability
        prediction = "AI-Generated" if fake_percentage > 0.5 else "Human-Written"
        
        # Format response for frontend
        response_data = {
            "ai_score": ai_score,  # 0.0 to 1.0 (e.g., 0.4459)
            "prediction": prediction,
            "details": {
                "confidence": f"{fake_percentage * 100:.1f}%",  # e.g., "44.6%"
                "human_confidence": f"{is_human * 100:.1f}%",  # e.g., "62.5%"
                "model_name": "RapidAPI AI Detector"
            }
        }
        
        return jsonify(response_data)
    
    except Exception as e:
        logging.error(f"Error in ai_refined: {str(e)}", exc_info=True)
        return jsonify({"error": "Internal server error"}), 500

# Plagiarism Detection Functions
WORD = re.compile(r'\w+')

def get_cosine(vec1, vec2):
    intersection = set(vec1.keys()) & set(vec2.keys())
    matchWords = {}
    for i in intersection:
        if vec1[i] > vec2[i]:
            matchWords[i] = vec2[i]
        else:
            matchWords[i] = vec1[i]
    numerator = sum([vec1[x] * matchWords[x] for x in intersection])
    sum1 = sum([vec1[x]**2 for x in vec1.keys()])
    sum2 = sum([matchWords[x]**2 for x in matchWords.keys()])
    denominator = math.sqrt(sum1) * math.sqrt(sum2)
    if denominator == 0:
        return 0.0
    else:
        return float(numerator) / denominator

def text_to_vector(text):
    words = WORD.findall(text)
    return Counter(words)

def cosineSim(text1, text2):
    t1 = text1.lower()
    t2 = text2.lower()
    vector1 = text_to_vector(t1)
    vector2 = text_to_vector(t2)
    cosine = get_cosine(vector1, vector2)
    return cosine

def getQueries(text, n):
    sentenceEnders = re.compile("['.!?]")
    sentenceList = sentenceEnders.split(text)
    sentencesplits = []
    en_stops = set(stopwords.words('english'))
    for sentence in sentenceList:
        x = re.compile(r'\W+', re.UNICODE).split(sentence)
        x = [word for word in x if word.lower() not in en_stops and word != '']
        sentencesplits.append(x)
    finalq = []
    for sentence in sentencesplits:
        l = len(sentence)
        if l > n:
            l = int(l/n)
            index = 0
            for i in range(0, l):
                finalq.append(sentence[index:index+n])
                index = index + n-1
                if index+n > l:
                    index = l-n-1
            if index != len(sentence):
                finalq.append(sentence[len(sentence)-index:len(sentence)])
        else:
            if l > 4:
                finalq.append(sentence)
    return finalq

def searchWeb(text, output, c):
    try:
        resource = build("customsearch", 'v1', developerKey=searchEngine_API).cse()
        result = resource.list(q=text, cx=searchEngine_Id).execute()
        searchInfo = result['searchInformation']
        if int(searchInfo['totalResults']) > 0:
            maxSim = 0
            itemLink = ''
            numList = len(result['items'])
            if numList >= 5:
                numList = 5
            for i in range(0, numList):
                item = result['items'][i]
                content = item['snippet']
                simValue = cosineSim(text, content)
                if simValue > maxSim:
                    maxSim = simValue
                    itemLink = item['link']
                if item['link'] in output:
                    itemLink = item['link']
                    break
            if itemLink in output:
                print('if', maxSim)
                output[itemLink] = output[itemLink] + 1
                c[itemLink] = ((c[itemLink] * (output[itemLink]-1) + maxSim)/(output[itemLink]))
            else:
                print('else', maxSim)
                print(text)
                print(itemLink)
                output[itemLink] = 1
                c[itemLink] = maxSim
    except Exception as e:
        print(text)
        print(e)
        print('error')
        return output, c, 1
    return output, c, 0

def findSimilarity(text):
    n = 9
    queries = getQueries(text, n)
    print('GetQueries task complete')
    q = [' '.join(d) for d in queries]
    output = {}
    c = {}
    while "" in q:
        q.remove("")
    count = len(q)
    if count > 100:
        count = 100
    numqueries = count
    for s in q[0:count]:
        output, c, errorCount = searchWeb(s, output, c)
        print('Web search task complete')
        numqueries = numqueries - errorCount
    totalPercent = 0
    outputLink = {}
    for link in output:
        percentage = (output[link]*c[link]*100)/numqueries
        if percentage > 10:
            totalPercent = totalPercent + percentage
            prevlink = link
            outputLink[link] = percentage
        elif 'prevlink' in locals() and len(prevlink) != 0:
            totalPercent = totalPercent + percentage
            outputLink[prevlink] = outputLink[prevlink] + percentage
        elif c[link] == 1:
            totalPercent = totalPercent + percentage
        print(link, totalPercent)
    print(count, numqueries)
    print(totalPercent, outputLink)
    print("\nDone!")
    return totalPercent, outputLink

def extract_text_from_txt(file):
    return file.read().decode('utf-8')

def extract_text_from_docx(file):
    doc = Document(file)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text

def extract_text_from_pdf(file_path):
    text = ""
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text()
    except Exception as e:
        print(f"Error extracting text from PDF: {e}")
    return text

def process_file(file):
    filename = file.filename
    text = ""
    if filename.endswith('.txt'):
        text = extract_text_from_txt(file)
    elif filename.endswith('.docx'):
        text = extract_text_from_docx(file)
    elif filename.endswith('.pdf'):
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)
        text = extract_text_from_pdf(file_path)
        os.remove(file_path)
    return text

@app.route('/test', methods=['POST'])
def test():
    print("Plagiarism test with text input")
    if 'q' in request.form and request.form['q']:
        percent, link = findSimilarity(request.form['q'])
        percent = round(percent, 2)
        return render_template('plagiarism_checker.html', link=link, percent=percent)
    return redirect(url_for('plagiarism_checker'))

@app.route('/file_test', methods=['POST'])
def file_test():
    print("Plagiarism test with file input")
    if 'docfile' in request.files:
        file = request.files['docfile']
        if file.filename != '':
            text = process_file(file)
            if text:
                percent, link = findSimilarity(text)
                percent = round(percent, 2)
                return render_template('plagiarism_checker.html', link=link, percent=percent)
    return redirect(url_for('plagiarism_checker'))

# Code Clustering
device = 0 if torch.cuda.is_available() else -1
feature_extractor = pipeline(
    "feature-extraction",
    model="microsoft/graphcodebert-base",
    tokenizer="microsoft/graphcodebert-base",
    device=device
)

def extract_embeddings(code_snippets):
    try:
        logging.info(f"Extracting embeddings for {len(code_snippets)} snippets")
        embeddings = feature_extractor(code_snippets, truncation=True, padding=True, max_length=512)
        if not embeddings or len(embeddings) == 0:
            logging.error("No embeddings extracted. Check input data or model loading.")
            return np.array([])
        cls_embeddings = np.array([np.array(embed).squeeze()[0] for embed in embeddings])
        logging.info(f"Extracted {cls_embeddings.shape[0]} embeddings with shape {cls_embeddings.shape}")
        return cls_embeddings
    except Exception as e:
        logging.error(f"Error extracting embeddings: {str(e)}", exc_info=True)
        return np.array([])

def cluster_code_files(files, threshold=0.875):
    try:
        if not files or not all(isinstance(f, dict) for f in files):
            logging.error("Invalid file format received for clustering.")
            return []
        texts = [f.get("content", "").strip() for f in files]
        names = [f.get("name", "") for f in files]
        if not any(texts):
            logging.warning("No valid file contents found.")
            return []
        if len(texts) == 1:
            logging.info("Only one file provided, creating a single cluster without similarity.")
            return [{
                "cluster_id": 0,
                "names": names,
            }]
        embeddings = extract_embeddings(texts)
        if embeddings.size == 0:
            logging.error("Embeddings are empty. Clustering cannot proceed.")
            return []
        embeddings = normalize(embeddings, axis=1)
        cosine_sim = cosine_similarity(embeddings)
        logging.info(f"Cosine Similarity Matrix:\n{cosine_sim}")
        clusters = []
        assigned = set()
        for i, name in enumerate(names):
            if name in assigned:
                continue
            new_cluster = {name}
            assigned.add(name)
            for j in range(len(names)):
                if i != j and names[j] not in assigned and cosine_sim[i, j] > threshold:
                    new_cluster.add(names[j])
                    assigned.add(names[j])
            clusters.append(list(new_cluster))
        cluster_list = []
        for cluster_id, filenames in enumerate(clusters):
            indices = [names.index(f) for f in filenames]
            if len(filenames) == 1:
                cluster_list.append({
                    "cluster_id": cluster_id,
                    "names": filenames,
                })
            else:
                avg_similarity = round(np.mean(cosine_sim[np.ix_(indices, indices)]), 2) * 100
                cluster_list.append({
                    "cluster_id": cluster_id,
                     "names": filenames,
                    "similarity": avg_similarity
                })
        logging.info(f"Clusters JSON: {json.dumps(cluster_list, indent=4)}")
        return cluster_list
    except Exception as e:
        logging.error(f"Clustering error: {str(e)}", exc_info=True)
        return []

@app.route('/cluster_python_folder/', methods=['POST'])
def cluster_python_folder():
    try:
        data = request.get_json()
        logging.info(f"Received Data: {json.dumps(data, indent=2)}")
        if not data or "files" not in data:
            return jsonify({"error": "Invalid request format"}), 400
        files_data = data["files"]
        if not isinstance(files_data, list):
            return jsonify({"error": "Files must be a list"}), 400
        clusters = cluster_code_files(files_data)
        if not clusters:
            logging.warning("No clusters were formed. Verify threshold or embeddings.")
        logging.info(f"Clusters formed: {clusters}")
        return jsonify({"clusters": clusters}), 200
    except Exception as e:
        logging.error(f"API error: {str(e)}", exc_info=True)
        return jsonify({"error": "Internal server error"}), 500

@app.route('/generate_cluster_pdf', methods=['POST'])
def generate_cluster_pdf():
    try:
        data = request.get_json()
        clusters = data.get('clusters', [])
        html = render_template_string('''
            <!DOCTYPE html>
            <html>
            <head>
                <style>
                    body { font-family: Arial, sans-serif; margin: 1in; line-height: 1.6; }
                    h1 { color: #2b2d42; font-size: 24px; border-bottom: 2px solid #4361ee; padding-bottom: 5px; }
                    .report-info { font-size: 12px; color: #666; margin-bottom: 20px; }
                    .cluster { margin: 20px 0; padding: 10px; background-color: #f9f9f9; border-radius: 5px; }
                    .cluster-header { display: flex; justify-content: space-between; align-items: center; margin-bottom: 10px; }
                    .cluster-title { font-size: 18px; font-weight: bold; color: #2b2d42; }
                    .similarity { color: #4361ee; font-weight: bold; }
                    ul { list-style-type: none; padding-left: 20px; }
                    li { margin: 5px 0; font-size: 12px; }
                    li:before { content: "• "; color: #4361ee; }
                </style>
            </head>
            <body>
                <h1>Code Cluster Analysis Report</h1>
                <div class="report-info">Generated on: {{ timestamp }}</div>
                {% if clusters %}
                    {% for cluster in clusters %}
                    <div class="cluster">
                        <div class="cluster-header">
                            <span class="cluster-title">Cluster {{ cluster.cluster_id }}</span>
                            {% if cluster.similarity is defined %}
                                <span class="similarity">Similarity: {{ cluster.similarity }}%</span>
                            {% endif %}
                        </div>
                        <ul>
                            {% for file in cluster.names %}
                            <li>{{ file }}</li>
                            {% endfor %}
                        </ul>
                    </div>
                    {% endfor %}
                {% else %}
                    <p>No clusters found.</p>
                {% endif %}
            </body>
            </html>
        ''', clusters=clusters, timestamp=datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
        try:
            config = pdfkit.configuration()
        except Exception as e:
            logging.warning(f"wkhtmltopdf not found in PATH, using hardcoded path: {str(e)}")
            config = pdfkit.configuration(wkhtmltopdf='C:/Program Files/wkhtmltopdf/bin/wkhtmltopdf.exe')
        pdf = pdfkit.from_string(html, False, options={
            'page-size': 'A4',
            'margin-top': '0.5in',
            'margin-right': '0.5in',
            'margin-bottom': '0.5in',
            'margin-left': '0.5in',
            'encoding': 'UTF-8'
        }, configuration=config)
        response = make_response(pdf)
        response.headers['Content-Type'] = 'application/pdf'
        response.headers['Content-Disposition'] = 'attachment; filename=cluster_report.pdf'
        return response
    except Exception as e:
        logging.error(f"PDF generation error: {str(e)}")
        return jsonify({'error': f'Failed to generate PDF: {str(e)}'}), 500

# SQLite Database Functions
def create_assignment_table(assignment_name):
    try:
        table_name = re.sub(r'[^a-zA-Z0-9_]', '_', assignment_name)
        if not table_name or table_name[0].isdigit():
            table_name = f"assignment_{table_name}"
        conn = sqlite3.connect('assignments.db')
        cursor = conn.cursor()
        cursor.execute(f'''
            CREATE TABLE IF NOT EXISTS {table_name} (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                file_name TEXT NOT NULL,
                ai_percentage REAL NOT NULL,
                plagiarism_percentage REAL NOT NULL
            )
        ''')
        conn.commit()
        conn.close()
        logging.info(f"Table {table_name} created or already exists.")
        return table_name
    except Exception as e:
        logging.error(f"Error creating table: {str(e)}")
        return None

def insert_analysis_results(table_name, file_name, ai_percentage, plagiarism_percentage):
    try:
        conn = sqlite3.connect('assignments.db')
        cursor = conn.cursor()
        cursor.execute(f'''
            INSERT INTO {table_name} (file_name, ai_percentage, plagiarism_percentage)
            VALUES (?, ?, ?)
        ''', (file_name, ai_percentage, plagiarism_percentage))
        conn.commit()
        conn.close()
        logging.info(f"Inserted {file_name} into {table_name}.")
    except Exception as e:
        logging.error(f"Error inserting data: {str(e)}")

def get_analysis_results(table_name):
    try:
        conn = sqlite3.connect('assignments.db')
        cursor = conn.cursor()
        cursor.execute(f'SELECT id, file_name, ai_percentage, plagiarism_percentage FROM {table_name}')
        results = cursor.fetchall()
        conn.close()
        return [{"id": row[0], "file_name": row[1], "ai_percentage": row[2], "plagiarism_percentage": row[3]} for row in results]
    except Exception as e:
        logging.error(f"Error retrieving data: {str(e)}")
        return []

def get_all_assignment_tables():
    try:
        conn = sqlite3.connect('assignments.db')
        cursor = conn.cursor()
        cursor.execute("SELECT name FROM sqlite_master WHERE type='table' AND name LIKE 'assignment_%'")
        tables = cursor.fetchall()
        conn.close()
        return [table[0] for table in tables]
    except Exception as e:
        logging.error(f"Error retrieving assignment tables: {str(e)}")
        return []
    
def process_assignment_files(files):
    """
    Process files for assignment upload, supporting multiple file types and folders
    
    Args:
        files: List of uploaded files or files in a folder
    
    Returns:
        List of processed file dictionaries with name and content
    """
    processed_files = []
    
    for file in files:
        # Handle individual file uploads
        if hasattr(file, 'filename'):
            filename = file.filename
            
            # Save temporary file if needed
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(file_path)
            
            try:
                # Text extraction based on file type
                if filename.endswith('.txt'):
                    with open(file_path, 'r', encoding='utf-8') as f:
                        text = f.read()
                elif filename.endswith('.docx'):
                    doc = Document(file_path)
                    text = "\n".join([para.text for para in doc.paragraphs])
                elif filename.endswith('.pdf'):
                    with open(file_path, 'rb') as f:
                        pdf_reader = PyPDF2.PdfReader(f)
                        text = "\n".join([page.extract_text() for page in pdf_reader.pages])
                else:
                    logging.warning(f"Unsupported file type: {filename}")
                    continue
                
                processed_files.append({
                    "name": filename,
                    "content": text
                })
            
            except Exception as e:
                logging.error(f"Error processing file {filename}: {str(e)}")
            
            finally:
                # Remove temporary file
                if os.path.exists(file_path):
                    os.remove(file_path)
        
        # Handle folder uploads (assuming file is actually a path string)
        elif isinstance(file, str) and os.path.isdir(file):
            for root, _, files in os.walk(file):
                for filename in files:
                    filepath = os.path.join(root, filename)
                    
                    # Process only txt, docx, and pdf files
                    if filename.endswith(('.txt', '.docx', '.pdf')):
                        try:
                            if filename.endswith('.txt'):
                                with open(filepath, 'r', encoding='utf-8') as f:
                                    text = f.read()
                            elif filename.endswith('.docx'):
                                doc = Document(filepath)
                                text = "\n".join([para.text for para in doc.paragraphs])
                            elif filename.endswith('.pdf'):
                                with open(filepath, 'rb') as f:
                                    pdf_reader = PyPDF2.PdfReader(f)
                                    text = "\n".join([page.extract_text() for page in pdf_reader.pages])
                            
                            processed_files.append({
                                "name": filename,
                                "content": text
                            })
                        
                        except Exception as e:
                            logging.error(f"Error processing file {filename}: {str(e)}")
    
    return processed_files

@app.route('/assignment_func', methods=['GET', 'POST'])
def assignment_func():
    if request.method == 'POST':
        data = request.get_json()
        assignment_name = data.get("assignment_name", "")
        files_data = data.get("files", [])
        if not assignment_name or not files_data:
            return jsonify({'error': 'Assignment name and files are required'}), 400
        
        table_name = create_assignment_table(assignment_name)
        if not table_name:
            return jsonify({'error': 'Failed to create assignment table'}), 500
        
        results = []
        for file in files_data:
            file_name = file.get("name", "")
            content = file.get("content", "")
            if not file_name or not content:
                logging.warning(f"Skipping {file_name}: No name or content provided")
                continue
            
            # Handle file content based on type
            if file_name.endswith('.txt'):
                text = content  # Already read as text by FileReader
            elif file_name.endswith('.pdf'):
                try:
                    if isinstance(content, str) and content.startswith('data:application/pdf;base64,'):
                        base64_string = content.split(',')[1]
                        pdf_bytes = base64.b64decode(base64_string)
                    else:
                        logging.error(f"Unexpected PDF content format for {file_name}: {content[:50]}")
                        continue
                    text = extract_text_from_pdf_buffer(pdf_bytes)
                except Exception as e:
                    logging.error(f"Failed to process PDF {file_name}: {str(e)}")
                    text = ""
            else:
                logging.warning(f"Unsupported file type: {file_name}")
                continue
            
            if not text:
                logging.warning(f"No text extracted from {file_name}")
                continue
            
            # AI Detection using ai_refined endpoint
            ai_response = requests.post(
                "http://127.0.0.1:5000/ai_refined/",
                json={"text": text}
            ).json()
            print("AI Detection Response:", ai_response)
            
            if "error" in ai_response:
                logging.error(f"AI detection failed for {file_name}: {ai_response['error']}")
                ai_percentage = 0.0
            else:
                ai_percentage = round(float(ai_response["ai_score"]) * 100, 2)
            
            # Plagiarism Detection
            try:
                plagiarism_percentage, _ = findSimilarity(text)
                plagiarism_percentage = round(plagiarism_percentage, 2)
            except Exception as e:
                logging.error(f"Plagiarism detection failed for {file_name}: {str(e)}")
                plagiarism_percentage = 0.0
            
            insert_analysis_results(table_name, file_name, ai_percentage, plagiarism_percentage)
            stored_results = get_analysis_results(table_name)
            results = stored_results
        
        return jsonify({
            "table_name": table_name,
            "results": results
        })
    
    elif request.method == 'GET':
        assignment_name = request.args.get("assignment_name", "")
        if assignment_name:
            table_name = re.sub(r'[^a-zA-Z0-9_]', '_', assignment_name)
            if not table_name or table_name[0].isdigit():
                table_name = f"assignment_{table_name}"
            results = get_analysis_results(table_name)
            return jsonify({
                "table_name": table_name,
                "results": results
            })
        else:
            tables = get_all_assignment_tables()
            all_assignments = {}
            for table in tables:
                results = get_analysis_results(table)
                all_assignments[table] = results
            return jsonify({
                "assignments": all_assignments
            })

@app.route('/generate_assignment_report', methods=['POST'])
def generate_assignment_report():
    try:
        data = request.get_json()
        assignment_name = data.get('assignment_name', '')
        results = data.get('results', [])
        
        html = render_template_string('''
            <!DOCTYPE html>
            <html>
            <head>
                <style>
                    body { font-family: Arial, sans-serif; margin: 1in; line-height: 1.6; }
                    h1 { color: #2b2d42; font-size: 24px; border-bottom: 2px solid #4361ee; padding-bottom: 5px; }
                    .report-info { font-size: 12px; color: #666; margin-bottom: 20px; }
                    table { width: 100%; border-collapse: collapse; margin: 20px 0; }
                    th, td { border: 1px solid #ddd; padding: 8px; text-align: left; }
                    th { background-color: #f2f2f2; color: #333; }
                    tr:nth-child(even) { background-color: #f9f9f9; }
                    .warning { background-color: #ffdddd; }
                    .caution { background-color: #ffffcc; }
                </style>
            </head>
            <body>
                <h1>Assignment Analysis Report: {{ assignment_name }}</h1>
                <div class="report-info">Generated on: {{ timestamp }}</div>
                
                <table>
                    <thead>
                        <tr>
                            <th>ID</th>
                            <th>File Name</th>
                            <th>AI-Generated Content (%)</th>
                            <th>Plagiarism (%)</th>
                        </tr>
                    </thead>
                    <tbody>
                        {% for result in results %}
                        <tr class="{% if result.ai_percentage > 70 or result.plagiarism_percentage > 30 %}warning{% elif result.ai_percentage > 50 or result.plagiarism_percentage > 20 %}caution{% endif %}">
                            <td>{{ result.id }}</td>
                            <td>{{ result.file_name }}</td>
                            <td>{{ "%.1f"|format(result.ai_percentage) }}%</td>
                            <td>{{ "%.1f"|format(result.plagiarism_percentage) }}%</td>
                        </tr>
                        {% endfor %}
                    </tbody>
                </table>
            </body>
            </html>
        ''', assignment_name=assignment_name, results=results, timestamp=datetime.now().strftime("%Y-%m-d %H:%M:%S"))
        
        try:
            config = pdfkit.configuration()
        except Exception as e:
            logging.warning(f"wkhtmltopdf not found in PATH, using hardcoded path: {str(e)}")
            config = pdfkit.configuration(wkhtmltopdf='C:/Program Files/wkhtmltopdf/bin/wkhtmltopdf.exe')
            
        pdf = pdfkit.from_string(html, False, options={
            'page-size': 'A4',
            'margin-top': '0.5in',
            'margin-right': '0.5in',
            'margin-bottom': '0.5in',
            'margin-left': '0.5in',
            'encoding': 'UTF-8'
        }, configuration=config)
        
        response = make_response(pdf)
        response.headers['Content-Type'] = 'application/pdf'
        response.headers['Content-Disposition'] = f'attachment; filename={re.sub(r"[^a-zA-Z0-9_]", "_", assignment_name)}_report.pdf'
        return response
    
    except Exception as e:
        logging.error(f"PDF generation error: {str(e)}")
        return jsonify({'error': f'Failed to generate PDF: {str(e)}'}), 500
# app.py (updated endpoint)
@app.route('/api/assignments')
def get_all_assignments():
    try:
        conn = sqlite3.connect('assignments.db')
        cursor = conn.cursor()
        
        # Get all non-system tables
        cursor.execute("SELECT name FROM sqlite_master WHERE type='table' AND name NOT LIKE 'sqlite_%'")
        tables = [row[0] for row in cursor.fetchall()]
        
        assignments = {}
        
        for table in tables:
            try:
                # Verify table structure
                cursor.execute(f"PRAGMA table_info({table})")
                columns = [col[1].lower() for col in cursor.fetchall()]
                required_columns = {'id', 'file_name', 'ai_percentage', 'plagiarism_percentage'}
                
                if not required_columns.issubset(columns):
                    logging.warning(f"Skipping table {table} - invalid structure")
                    continue
                
                # Get data
                cursor.execute(f'''
                    SELECT id, file_name, 
                           CAST(ai_percentage AS REAL), 
                           CAST(plagiarism_percentage AS REAL)
                    FROM {table}
                ''')
                results = cursor.fetchall()
                
                # Clean table name
                clean_name = table.replace('assignment_', ' ').replace('_', ' ').title()
                assignments[clean_name.strip()] = [
                    {
                        "id": row[0],
                        "file_name": row[1],
                        "ai_percentage": float(row[2]),
                        "plagiarism_percentage": float(row[3])
                    } for row in results
                ]
                
            except sqlite3.Error as e:
                logging.error(f"Error processing {table}: {str(e)}")
                continue
        
        conn.close()
        return jsonify({"assignments": assignments})
    
    except Exception as e:
        logging.error(f"Database error: {str(e)}")
        return jsonify({"error": "Server error"}), 500
if __name__ == '__main__':
    app.run(debug=True)